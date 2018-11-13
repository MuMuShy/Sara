#!/usr/bin/env python
# coding: utf-8

# In[24]:


from docx import Document
from docx.shared import Inches
import requests
from bs4 import BeautifulSoup
import re
import os
import time
def downloadImage(downloadname):
    from google_images_download import google_images_download   #importing the library
    response = google_images_download.googleimagesdownload()   #class instantiation
    path1=os.path.abspath('..')
    arguments = {"keywords":downloadname,"limit":1,"print_urls":True,"output_directory":path1,"image_directory":downloadname}   #creating list of arguments
    paths = response.download(arguments)   #passing the arguments to the function
    print('done')
def getImagefileName(downloadname):
    from os import listdir
    from os.path import isfile, isdir, join
    path1=os.path.abspath('..')
    # 指定要列出所有檔案的目錄
    mypath = path1+"/"+downloadname+"/"
    # 取得所有檔案與子目錄名稱
    files = listdir(mypath)
    # 以迴圈處理
    for f in files:
      # 產生檔案的絕對路徑
      fullpath = join(mypath, f)
      # 判斷 fullpath 是檔案還是目錄
      if isfile(fullpath):
        return(str(f))
      elif isdir(fullpath):
        print("目錄：", f)
def MakeReport():
    from SaraModules import SaraSpeak
    from SaraModules import SaraListen
    SaraSpeak.SpeakChinese('請問你要做什麼主題的報告')
    time.sleep(2)
    temp=SaraListen.Listen()
    if temp:
        Doctitle=temp
    else:
        SaraSpeak.SpeakChinese('不好意思我不太了解 你可以輸入給我嗎')
        Doctitle=input('請問你要做什麼報告?')
    downloadImage(Doctitle)
    path1=os.path.abspath('..')
    imagePath=path1+"/"+Doctitle+"/"+getImagefileName(Doctitle)
    print("相片路徑"+imagePath)
    document = Document()
    document.add_heading(Doctitle, 0)
    res=requests.get('https://zh.wikipedia.org/wiki/{}'.format(Doctitle))
    soup=BeautifulSoup(res.text,'lxml')
    result=soup.find_all('p')
    par1=str(result[0].text)
    par1 = re.sub(u"\\(.*?\\)|\\{.*?}|\\[.*?]", "", par1)
    document.add_paragraph(par1)
    document.add_picture(imagePath, width=Inches(6), height=Inches(4))
    result2=soup.find_all('h3')
    for r in result2:
        result=str(r.text)
        result = re.sub(u"\\(.*?\\)|\\{.*?}|\\[.*?]", "", result)
        print("標題 :"+result)
        para = r.find_next_sibling('p')
        if para==None:
            docpath=os.path.abspath('..')
            document.save(docpath+"/"+Doctitle+".doc")
            return
        else:       
            para=str(para.text)
            if len(para)<10:
                continue
            para=re.sub(u"\\(.*?\\)|\\{.*?}|\\[.*?]", "", para)
            document.add_heading(result, 0)
            document.add_paragraph(para)
            print("內容 : "+para)


# In[ ]:




